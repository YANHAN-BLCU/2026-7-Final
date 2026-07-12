from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTDIR = r"F:\资料\机器学习\真题\题解"
os.makedirs(OUTDIR, exist_ok=True)


def set_font(run, size=10.5, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text="", style=None, bold=False, center=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold)
    return p


def formula(doc, text):
    p = para(doc, center=True)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    r.font.size = Pt(11)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_font(r, size=15 if level == 1 else 12.5, bold=True, color="2E74B5")


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(str(h))
        set_font(r, size=9.5, bold=True)
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            r = cs[i].paragraphs[0].add_run(str(v))
            set_font(r, size=9.5)
    para(doc)


def setup(title, subtitle):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.85)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(10.5)
    p = para(doc, title, bold=True, center=True)
    for r in p.runs:
        set_font(r, size=20, bold=True, color="1F4D78")
    p = para(doc, subtitle, center=True)
    for r in p.runs:
        set_font(r, size=11, color="555555")
    return doc


def save(doc, name):
    path = os.path.join(OUTDIR, name)
    doc.save(path)
    return path


def make_2024():
    doc = setup("2024 年机器学习真题题解", "按题型整理，公式使用可稳定渲染的数学排版")
    heading(doc, "一、判断题")
    tbl(doc, ["题号", "答案", "要点"], [
        [1, "对", "按标注情况可分为有监督、半监督、无监督学习。"],
        [2, "错", "训练时不能直接最小化测试误差，通常最小化经验风险或结构风险。"],
        [3, "对", "学习算法需要归纳偏好；空间偏好可等价反映算法偏好。"],
        [4, "对", "ID3 逐层贪心选择属性，不能保证全局最优。"],
        [5, "对", "概念学习通常是有监督二分类问题。"],
        [6, "错", "kNN 是惰性学习，无感知机式收敛训练过程。"],
        [7, "对", "最大间隔要求正确分类且离分离超平面尽量远。"],
        [8, "错", "硬间隔支持向量在间隔边界；软间隔下可在间隔内或被误分。"],
        [9, "对", "逻辑回归估计后验概率，可用于分类。"],
        [10, "对", "均匀先验下，最大后验估计等价于最大似然估计。"],
    ])
    heading(doc, "二、支持向量机")
    formula(doc, "f(x)=w·x+b")
    formula(doc, "函数间隔：γ̂_i=y_i(w·x_i+b)")
    formula(doc, "几何间隔：γ_i=y_i(w·x_i+b)/||w||")
    para(doc, "函数间隔受参数缩放影响，几何间隔消除尺度影响。若 y_i(w·x_i+b)>0，则样本被正确分类。")
    formula(doc, "L_hinge(y,f(x))=max(0,1-yf(x))")
    formula(doc, "min_{w,b} 1/2||w||² + C∑_{i=1}^N max(0,1-y_i(w·x_i+b))")
    para(doc, "当 y_i(w·x_i+b)<1 时产生合页损失；当 y_i(w·x_i+b)≥1 时损失为 0。支持向量是决定最优超平面的样本。")
    heading(doc, "三、朴素贝叶斯")
    formula(doc, "P(x|y)=∏_{j=1}^d P(a_j|y)")
    formula(doc, "h(x)=argmax_{y∈Y} P(y)∏_{j=1}^d P(a_j|y)")
    formula(doc, "m-估计：P(a_j=v|y=c)=(N_{c,j,v}+mp)/(N_c+m)")
    formula(doc, "P(Yes|x)≈0.00327,  P(No|x)≈0.00581")
    para(doc, "因此对 <sunny,cool,high,strong> 的预测为 No，即不进行。")
    heading(doc, "四、决策树")
    formula(doc, "H(D)=-∑_k p_k log₂p_k")
    formula(doc, "g(D,A)=H(D)-H(D|A)")
    formula(doc, "g_R(D,A)=g(D,A)/H_A(D)")
    tbl(doc, ["属性", "信息增益"], [["天气", "0.247"], ["温度", "0.029"], ["湿度", "0.048"], ["风速", "0.048"]])
    para(doc, "根节点选择“天气”。最终树：天气=阴→进行；天气=晴时按湿度划分，高→取消，正常→进行；天气=雨时按风速划分，弱→进行，强→取消。")
    heading(doc, "五、AdaBoost")
    formula(doc, "α_m=1/2 ln((1-e_m)/e_m)")
    formula(doc, "w_{m+1,i}=w_{m,i} exp(-α_m y_iG_m(x_i))/Z_m")
    para(doc, "第一轮可取“身体”阈值 v=1，误差 e1=0.4，α1≈0.203。第一轮后权重最大样本为 3、4、5、9，权重均为 0.125。")
    para(doc, "第二轮可取“业务”阈值 v=3，误差 e2=0.375，α2≈0.255。第二轮后权重最大样本为 9，权重为 1/6≈0.167。")
    formula(doc, "G(x)=sign(0.203G_1(x)+0.255G_2(x))")
    return save(doc, "2024年机器学习真题题解.docx")


def make_2025():
    doc = setup("2025 年机器学习真题题解", "按题型整理，公式使用可稳定渲染的数学排版")
    heading(doc, "一、判断题")
    tbl(doc, ["题号", "答案", "要点"], [
        [1, "对", "训练集和测试集通常假设同分布。"],
        [2, "对", "学习目标是降低真实风险或测试误差。"],
        [3, "错", "最大间隔并非因为损失函数与几何间隔等价。"],
        [4, "对", "机器学习算法一般都有归纳偏好。"],
        [5, "对", "决策树偏好纯度提升强的划分属性。"],
        [6, "对", "剪枝可缓解过拟合。"],
        [7, "对", "感知机线性可分收敛，线性不可分不保证收敛。"],
        [8, "错", "朴素贝叶斯假设是类别条件下属性独立。"],
        [9, "对", "Bagging 相对独立，Boosting 串行依赖。"],
        [10, "对", "均匀先验下 MAP 等价于 MLE。"],
    ])
    heading(doc, "二、感知机")
    formula(doc, "f(x)=sign(w·x+b)")
    formula(doc, "原始目标：min L(w,b)=-∑_{x_i∈M} y_i(w·x_i+b)")
    formula(doc, "对偶形式：w=∑_{i=1}^N α_i y_i x_i")
    formula(doc, "若 y_i(w·x_i+b)≤0：w←w+ηy_ix_i, b←b+ηy_i")
    para(doc, "数据线性不可分时，普通感知机可能不收敛，可设置最大迭代次数、采用口袋算法、核方法或改用软间隔分类器。")
    heading(doc, "三、支持向量机")
    formula(doc, "γ_i=y_i(w·x_i+b)/||w||")
    formula(doc, "硬间隔：min 1/2||w||²,  s.t. y_i(w·x_i+b)≥1")
    formula(doc, "软间隔：min 1/2||w||²+C∑ξ_i,  s.t. y_i(w·x_i+b)≥1-ξ_i")
    para(doc, "支持向量是拉格朗日乘子非零、决定分离超平面的样本。线性不可分时可用软间隔或核函数处理。")
    heading(doc, "四、决策树")
    formula(doc, "g_R(D,A)=g(D,A)/H_A(D)")
    tbl(doc, ["属性", "信息增益", "划分熵", "信息增益比"], [
        ["年龄", "0.083", "1.585", "0.052"],
        ["有工作", "0.324", "0.918", "0.352"],
        ["有自己的房子", "0.420", "0.971", "0.433"],
        ["信贷情况", "0.363", "1.566", "0.232"],
    ])
    para(doc, "根节点为“有自己的房子”。有房子→类别=是；无房子时按“有工作”划分，有工作→是，无工作→否。")
    heading(doc, "五、AdaBoost")
    para(doc, "集成学习把多个弱学习器组合为强学习器。主要分为 Bagging 和 Boosting：Bagging 并行训练以降低方差，Boosting 串行训练并提高难分样本权重。")
    formula(doc, "e_m=∑_{i=1}^N w_{m,i} I(G_m(x_i)≠y_i)")
    formula(doc, "α_m=1/2 ln((1-e_m)/e_m)")
    formula(doc, "w_{m+1,i}=w_{m,i} exp(-α_m y_iG_m(x_i))/Z_m")
    para(doc, "Z_m 为归一化因子。误分样本权重上升，正确分类样本权重下降。")
    return save(doc, "2025年机器学习真题题解.docx")


if __name__ == "__main__":
    print(make_2024())
    print(make_2025())
