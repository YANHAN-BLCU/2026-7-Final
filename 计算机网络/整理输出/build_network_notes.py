from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
OUT_DIR = BASE / "整理输出"
QUESTION_TXT = OUT_DIR / "计算机网络题库_文本提取.txt"
DOCX_OUT = OUT_DIR / "计算机网络串讲与题库精讲.docx"
MD_INDEX_OUT = OUT_DIR / "题库100题章节索引.md"


CHAPTERS = {
    "第1章 概述": {
        "outline": [
            "计算机网络的定义、组成、分类与性能指标",
            "互联网边缘部分与核心部分：端系统、通信方式、分组交换",
            "计算机网络体系结构：协议、接口、服务、OSI、TCP/IP、五层模型",
        ],
        "thread": "先建立“网络如何分层工作”的总框架，再把时延、交换方式和协议三要素作为后续各层题目的共同语言。",
        "core": [
            "体系结构只规定层次、功能和协议，不考协议内部实现细节。",
            "OSI 自下而上：物理、数据链路、网络、运输、会话、表示、应用；TCP/IP 常按网络接口、网际、运输、应用理解。",
            "相邻节点通信、主机到主机通信、端到端进程通信要分清：链路层、网络层、运输层各管一段。",
            "分组交换时延常由发送时延、传播时延、排队时延、处理时延组成；存储转发题要按链路段逐段画时间线。",
        ],
        "pitfalls": [
            "把网络层的路由选择误认为运输层功能。",
            "把会话层、表示层、应用层功能混在一起：会话管对话，表示管格式/编码/加密压缩，应用直接面向应用进程。",
            "传输效率题漏算封装开销，或把物理层、应用层也误计入题设开销。",
        ],
        "examples": [1, 2, 3, 4, 5, 6, 13],
    },
    "第2章 物理层": {
        "outline": [
            "物理层基本概念与接口特性：机械、电气、功能、过程",
            "数据通信模型、信道基本概念、编码与调制",
            "信道极限容量：奈奎斯特定理与香农定理",
            "传输媒体、信道复用、数字传输系统与宽带接入",
        ],
        "thread": "物理层题目的核心是“信号如何承载比特”，计算题围绕码元、波特率、比特率、带宽和信噪比展开。",
        "core": [
            "码元速率是每秒发送多少个信号单元，单位 Baud；数据率是每秒传多少 bit，单位 bit/s。",
            "若一个码元有 V 种离散状态，则每个码元携带 log2(V) bit。",
            "无噪声信道最高码元速率常按 2W 理解，理想低通信道最高数据率为 2Wlog2(V)。",
            "有噪声信道极限信息率 C = W log2(1 + S/N)，注意 dB 与普通信噪比的换算。",
            "曼彻斯特编码每个比特中间跳变，差分曼彻斯特主要看比特开始处是否跳变。",
        ],
        "pitfalls": [
            "把传播速率当成数据传输速率。传播速率影响传播时延，不决定比特发送速率。",
            "把码元速率和比特率混用，忘记乘 log2(V)。",
            "看到“4 个幅值/4 种相位/16 种状态”时不先换算每码元比特数。",
        ],
        "examples": [16, 18, 19, 20, 21, 23, 27, 28, 31],
    },
    "第3章 数据链路层": {
        "outline": [
            "点对点信道的数据链路层：帧、封装成帧、透明传输、差错检测",
            "流量控制与可靠传输：停止-等待、GBN、SR",
            "PPP 协议",
            "广播信道与介质访问控制：ALOHA、CSMA、CSMA/CD、CSMA/CA",
            "以太网、集线器、交换机、扩展以太网与高速以太网",
        ],
        "thread": "数据链路层最容易出题，因为它把“相邻节点可靠传输”和“共享信道如何抢占”合在一起。题库高频集中在随机访问协议、以太网交换机和可靠传输窗口。",
        "core": [
            "封装成帧解决边界识别；透明传输解决数据中出现帧定界符；差错检测常见 CRC。",
            "停止-等待一次只允许一个未确认帧；GBN 累计确认但出错后回退重传；SR 对单个出错帧选择重传。",
            "ALOHA 不监听直接发；CSMA 先监听；CSMA/CD 用于有线以太网检测冲突；CSMA/CA 用于无线局域网尽量避免冲突。",
            "集线器工作在物理层，扩大冲突域；交换机工作在数据链路层，按 MAC 地址学习和转发，隔离冲突域。",
            "以太网最短帧长与争用期有关，本质是保证发送站能在发送结束前检测到冲突。",
        ],
        "pitfalls": [
            "把 CSMA/CD 用到无线网络，或把 CSMA/CA 用到传统有线以太网。",
            "交换机、集线器、路由器层次混淆：物理层、链路层、网络层。",
            "GBN 与 SR 的接收窗口、重传对象和确认方式混淆。",
        ],
        "examples": [32, 34, 36, 40, 43, 48, 50, 54, 56, 60, 62],
    },
    "第4章 网络层": {
        "outline": [
            "网络层服务模型：无连接、尽最大努力交付的数据报服务",
            "IPv4 地址、子网划分、CIDR 与路由聚合",
            "ARP、IP 分组、NAT、ICMP",
            "路由表与分组转发，最长前缀匹配",
            "路由选择协议：RIP、OSPF、BGP",
            "路由器组成与异构网络互联",
        ],
        "thread": "网络层把“从源主机到目的主机”落到 IP 地址、路由表和转发规则上。题库主要考地址计算、最长前缀匹配和路由协议性质。",
        "core": [
            "IP 层向上提供无连接、不可靠、尽最大努力的数据报服务。",
            "子网划分先看前缀长度，再算网络地址、广播地址、可用主机数和地址范围。",
            "CIDR 路由聚合用共同前缀表达多个连续网络；转发时优先最长前缀匹配。",
            "ARP 把 IP 地址解析为同一链路上的 MAC 地址。",
            "RIP 是距离向量协议，度量是跳数，16 表示不可达；OSPF 是链路状态协议；BGP 用于自治系统之间。",
        ],
        "pitfalls": [
            "看到默认路由就直接选，忘记最长前缀匹配优先。",
            "/30 网络可用主机数是 2，不是 4。",
            "把 ARP 理解为跨网络找 MAC；ARP 只解决同一链路下一跳的硬件地址。",
        ],
        "examples": [63, 64, 66, 67, 69, 70, 71, 72, 74, 78, 80],
    },
    "第5章 运输层": {
        "outline": [
            "运输层服务：进程到进程、端口、复用与分用",
            "UDP：无连接、首部简单、尽最大努力",
            "TCP：面向连接、可靠传输、流量控制、拥塞控制",
            "TCP 报文段格式、序号与确认号",
            "TCP 连接建立与释放，拥塞窗口变化",
        ],
        "thread": "运输层题目要把“可靠传输机制”和“拥塞控制曲线”分开看：前者关注字节序号、确认号、窗口；后者关注 cwnd、ssthresh、慢开始、拥塞避免、快重传快恢复。",
        "core": [
            "TCP 序号表示本报文段数据部分第一个字节的编号；确认号表示期望收到的下一个字节编号。",
            "三次握手解决双方收发能力确认和旧连接请求干扰；四次挥手源于 TCP 全双工连接的两个方向分别关闭。",
            "慢开始阶段 cwnd 指数增长，拥塞避免阶段线性增长；超时通常使 ssthresh 减半、cwnd 回到 1 MSS。",
            "收到 3 个重复 ACK 常触发快重传，随后进入快恢复，行为与超时不同。",
            "UDP 首部简单，无连接、不保证可靠交付，适合实时性或应用自控可靠性的场景。",
        ],
        "pitfalls": [
            "把确认号理解为“最后一个收到的字节号”，实际是“下一个期望字节号”。",
            "拥塞控制和流量控制混淆：拥塞控制面向网络负载，流量控制面向接收方能力。",
            "忽略 TCP 按字节编号，不是按报文段编号。",
        ],
        "examples": [82, 83, 84, 85, 88, 90, 91, 94, 95, 96],
    },
    "第6章 应用层": {
        "outline": [
            "应用层协议的工作方式：客户/服务器与应用进程通信",
            "DNS、HTTP、FTP、电子邮件、WWW 等常见协议",
            "电子邮件系统：用户代理、邮件服务器、SMTP、POP3、IMAP",
            "应用层协议与下层服务的关系",
        ],
        "thread": "应用层题目常把具体协议和下层服务绑定在一起考。题库最后几题集中在电子邮件，以及 PPP、ARP、UDP、WWW 等跨层协议识别。",
        "core": [
            "SMTP 用于邮件发送/转发，POP3 和 IMAP 用于用户读取邮件。",
            "POP3 建立在 TCP 连接上，使用可靠传输服务。",
            "DNS 负责域名到 IP 地址等资源记录的解析；HTTP 支撑万维网对象请求与响应。",
            "应用层协议不等于应用程序，协议规定通信规则，程序实现这些规则。",
        ],
        "pitfalls": [
            "把 SMTP、POP3、IMAP 的方向和用途混淆。",
            "只记协议名，不记它依赖 TCP 还是 UDP。",
            "跨层题要先定位协议所在层，再判断功能。",
        ],
        "examples": [97, 98, 99, 100],
    },
}


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clean_text(value: str) -> str:
    value = value.replace("\x0b", "\n").replace("\x01", "[原题图示]")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def parse_questions() -> list[dict]:
    text = QUESTION_TXT.read_text(encoding="utf-8", errors="ignore")
    text = text.replace("\x0b", "\n").replace("\x01", "[原题图示]")
    matches = list(re.finditer(r"(?:^|\n)\s*(\d+)、", text, flags=re.M))
    items = []
    for idx, match in enumerate(matches):
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        block = text[match.start() : end].strip()
        qno = int(match.group(1))
        answer = re.search(r"答案：\s*([A-D])", block)
        parse = re.search(r"解析：\s*(.*?)(?:\n\s*知识点：|$)", block, flags=re.S)
        kp = re.search(r"知识点：\s*(.*)", block, flags=re.S)

        lines = [line.strip() for line in block.splitlines() if line.strip()]
        stem_lines = []
        options = []
        in_options = False
        for line in lines:
            if re.match(r"^[A-D]、", line):
                in_options = True
            if line.startswith("答案："):
                break
            if in_options:
                options.append(line)
            else:
                stem_lines.append(line)

        items.append(
            {
                "no": qno,
                "stem": clean_text(" ".join(stem_lines)),
                "options": [clean_text(x) for x in options],
                "answer": answer.group(1) if answer else "",
                "analysis": clean_text(parse.group(1)) if parse else "",
                "kp": clean_text(kp.group(1)) if kp else "",
                "block": block,
            }
        )
    return items


def classify_question(item: dict) -> str:
    # The source question bank is ordered by the textbook/PPT chapters. Use that
    # ordering as the primary signal; keyword scoring below is only a fallback
    # for future inserts or out-of-range questions.
    ranges = [
        (1, 15, "第1章 概述"),
        (16, 30, "第2章 物理层"),
        (31, 62, "第3章 数据链路层"),
        (63, 80, "第4章 网络层"),
        (81, 96, "第5章 运输层"),
        (97, 100, "第6章 应用层"),
    ]
    for start, end, chapter in ranges:
        if start <= item["no"] <= end:
            return chapter

    keys = {
        "第1章 概述": ["分层", "OSI", "IOS", "TCP/IP", "协议", "接口", "服务", "性能指标", "分组交换", "报文交换", "电路交换", "时延", "体系结构"],
        "第2章 物理层": ["编码", "调制", "奈奎斯特", "香农", "码元", "波特", "速率", "物理层接口", "双绞线", "信道"],
        "第3章 数据链路层": ["ALOHA", "CSMA", "CSMA/CD", "CSMA/CA", "停止-等待", "GBN", "SR", "PPP", "检错", "纠错", "以太网", "交换机", "集线器", "802.11", "IEEE 802"],
        "第4章 网络层": ["CIDR", "IPv4", "IP", "NAT", "RIP", "OSPF", "BGP", "ARP", "ICMP", "子网", "路由", "网络层", "分组转发"],
        "第5章 运输层": ["TCP", "UDP", "拥塞", "连接管理", "TCP段", "传输层", "端口", "确认序号"],
        "第6章 应用层": ["电子邮件", "WWW", "DNS", "HTTP", "FTP", "SMTP", "POP3", "IMAP", "应用层"],
    }
    haystack = f"{item['kp']} {item['stem']}".lower()
    best = ("未归类", 0)
    for chapter, terms in keys.items():
        score = sum(1 for term in terms if term.lower() in haystack)
        if score > best[1]:
            best = (chapter, score)
    return best[0]


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 15, "1F4D78", 14, 7),
        ("Heading 2", 12.5, "2E74B5", 10, 5),
        ("Heading 3", 11, "1F4D78", 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("计算机网络串讲与题库精讲")
    set_run_font(run, size=8, color="666666")


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("计算机网络串讲与题库精讲")
    set_run_font(run, size=21, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("以课堂 PPT 为大纲，以《计算机网络题库.doc》为核心重点")
    set_run_font(run, size=10.5, color="555555")


def add_bullet(doc: Document, text: str, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=9.7)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=9.7)
    return p


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "F4F6F9")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    set_run_font(r, size=9.5, bold=True, color="1F4D78")
    r = p.add_run(body)
    set_run_font(r, size=9.5, color="222222")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(header_cells[i], "E8EEF5")
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def question_lookup(items: list[dict]) -> dict[int, dict]:
    return {item["no"]: item for item in items}


def compact_stem(item: dict, max_len=150) -> str:
    stem = item["stem"]
    stem = re.sub(r"^\d+、", "", stem).strip()
    if len(stem) > max_len:
        return stem[:max_len].rstrip() + "..."
    return stem


def add_question_explain(doc: Document, item: dict):
    doc.add_heading(f"题 {item['no']}｜{item['kp'] or '综合题'}", level=3)
    p = doc.add_paragraph()
    r = p.add_run("题干：")
    set_run_font(r, size=9.5, bold=True, color="1F4D78")
    r = p.add_run(compact_stem(item, 220))
    set_run_font(r, size=9.5)
    if item["options"] and len(item["options"]) <= 5:
        for option in item["options"]:
            add_bullet(doc, option, level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"答案：{item['answer'] or '原题未清晰提取'}。")
    set_run_font(r, size=9.5, bold=True, color="9B1C1C")
    analysis = item["analysis"] or "原题解析依赖图示或公式。复习时应先定位所属知识点，再按本章方法列式或判断协议功能。"
    p.add_run("讲解：")
    set_run_font(p.runs[-1], size=9.5, bold=True, color="1F4D78")
    r = p.add_run(analysis)
    set_run_font(r, size=9.5)


def make_markdown_index(items: list[dict]):
    grouped = {}
    for item in items:
        grouped.setdefault(item["chapter"], []).append(item)
    lines = ["# 计算机网络题库100题章节索引", ""]
    for chapter in CHAPTERS:
        rows = grouped.get(chapter, [])
        lines.append(f"## {chapter}（{len(rows)}题）")
        for item in rows:
            lines.append(f"- {item['no']:02d}. 答案 {item['answer'] or '-'}｜{item['kp']}｜{compact_stem(item, 80)}")
        lines.append("")
    if grouped.get("未归类"):
        lines.append("## 未归类")
        for item in grouped["未归类"]:
            lines.append(f"- {item['no']:02d}. 答案 {item['answer'] or '-'}｜{item['kp']}｜{compact_stem(item, 80)}")
    MD_INDEX_OUT.write_text("\n".join(lines), encoding="utf-8")


def build_doc(items: list[dict]):
    lookup = question_lookup(items)
    doc = Document()
    configure_document(doc)
    add_footer(doc)
    add_title(doc)

    doc.add_heading("使用方式", level=1)
    add_callout(
        doc,
        "复习路径",
        "先用每章“课件主线”建立框架，再背“基础知识串讲”和“易错点”，最后按“题库精讲”复盘同类题。不要孤立刷题，每道题都要回到层次、协议、公式或转发表。",
    )
    for text in [
        "第一遍：按第1章到第6章顺序读，重点理解每章的“层次职责”。",
        "第二遍：只看每章高频题型，把题库中的同类题归并成模板。",
        "第三遍：用附录索引回刷 100 题，错题旁边标注“错在概念、公式、协议层次、还是计算步骤”。",
    ]:
        add_number(doc, text)

    doc.add_heading("题库覆盖总览", level=1)
    counts = Counter(item["chapter"] for item in items)
    rows = []
    for chapter in CHAPTERS:
        data = CHAPTERS[chapter]
        rows.append([chapter, str(counts.get(chapter, 0)), "；".join(data["outline"][:2])])
    add_table(doc, ["章节", "题库题量", "课件主线"], rows, widths=[3.2, 2.0, 10.2])

    doc.add_heading("全书知识串联", level=1)
    for text in [
        "物理层回答“比特怎样变成信号在介质上传输”。关键词：带宽、码元、编码、调制、复用、信道容量。",
        "数据链路层回答“相邻节点怎样以帧为单位传输，并在共享信道上协调访问”。关键词：帧、差错检测、可靠传输、MAC、以太网、交换机。",
        "网络层回答“分组怎样跨多个网络到达目的主机”。关键词：IP 地址、子网、路由表、最长前缀匹配、路由协议、ARP。",
        "运输层回答“主机中的进程怎样可靠或尽力通信”。关键词：端口、UDP、TCP、序号、确认、窗口、拥塞控制、连接管理。",
        "应用层回答“具体网络应用怎样组织报文交互”。关键词：DNS、HTTP、FTP、电子邮件、SMTP、POP3、IMAP。",
    ]:
        add_bullet(doc, text)

    for idx, (chapter, data) in enumerate(CHAPTERS.items()):
        if idx > 0:
            doc.add_page_break()
        doc.add_heading(chapter, level=1)
        add_callout(doc, "课件主线", data["thread"])

        doc.add_heading("一、课堂 PPT 大纲", level=2)
        for item in data["outline"]:
            add_bullet(doc, item)

        doc.add_heading("二、基础知识串讲", level=2)
        for item in data["core"]:
            add_bullet(doc, item)

        doc.add_heading("三、题库高频考法", level=2)
        chapter_questions = [item for item in items if item["chapter"] == chapter]
        kp_counter = Counter()
        for item in chapter_questions:
            for part in re.split(r"[；;、，,]+", item["kp"]):
                part = part.strip()
                if part:
                    kp_counter[part] += 1
        if kp_counter:
            freq_rows = [[kp, str(count)] for kp, count in kp_counter.most_common(8)]
            add_table(doc, ["高频知识点", "出现次数"], freq_rows, widths=[10.8, 2.0])
        else:
            add_bullet(doc, "本章题目较少，复习时重点看概念边界和协议功能。")

        doc.add_heading("四、易错点提醒", level=2)
        for item in data["pitfalls"]:
            add_bullet(doc, item)

        doc.add_heading("五、题库典型题精讲", level=2)
        selected = [lookup[n] for n in data["examples"] if n in lookup]
        for item in selected:
            add_question_explain(doc, item)

    doc.add_page_break()
    doc.add_heading("附录：题库 100 题章节索引", level=1)
    doc.add_paragraph("下面索引用于回刷题库。题干较长或原题含图时，仅保留知识点和题干开头，完整题目请回到原题库。")
    for chapter in CHAPTERS:
        doc.add_heading(chapter, level=2)
        rows = []
        for item in items:
            if item["chapter"] == chapter:
                rows.append([str(item["no"]), item["answer"] or "-", item["kp"], compact_stem(item, 70)])
        add_table(doc, ["题号", "答案", "知识点", "题干提示"], rows, widths=[1.3, 1.3, 5.2, 8.0])

    doc.add_heading("最后复盘清单", level=1)
    for text in [
        "看到层次模型题：先判断问的是相邻节点、主机到主机，还是进程到进程。",
        "看到速率题：先分清码元速率、比特率、带宽、信噪比、传播速率。",
        "看到 MAC 访问题：先判断信道是有线以太网还是无线局域网。",
        "看到 IP 地址题：先写前缀，再算网络地址、广播地址和可用主机范围。",
        "看到 TCP 题：序号看“第一个字节”，确认号看“下一个期望字节”。",
        "看到应用层题：把协议用途、方向和下层 TCP/UDP 服务一起记。",
    ]:
        add_bullet(doc, text)

    doc.save(DOCX_OUT)


def main():
    items = parse_questions()
    for item in items:
        item["chapter"] = classify_question(item)
    make_markdown_index(items)
    build_doc(items)
    print(f"questions={len(items)}")
    print(f"docx={DOCX_OUT}")
    print(f"index={MD_INDEX_OUT}")
    for chapter, count in Counter(item["chapter"] for item in items).items():
        print(f"{chapter}: {count}")


if __name__ == "__main__":
    main()
