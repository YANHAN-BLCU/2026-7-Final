from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path(__file__).resolve().parent
FIG_DIR = OUT_DIR / "figures"
FIG_DIR.mkdir(exist_ok=True)

QUESTION_DOCX = OUT_DIR / "计算机组成原理期末高频必考模拟试卷A卷.docx"
ANSWER_DOCX = OUT_DIR / "计算机组成原理期末高频必考模拟试卷A卷_答案解析.docx"
SYNC_DOCX = OUT_DIR / "计算机组成原理期末高频必考模拟试卷B卷_提高补充卷.docx"
SYNC_ANSWER_DOCX = OUT_DIR / "计算机组成原理期末高频必考模拟试卷B卷_提高补充卷_答案解析.docx"
C2025_DOCX = OUT_DIR / "计算机组成原理期末高频必考模拟试卷C卷_2025真题风格补充.docx"
C2025_ANSWER_DOCX = OUT_DIR / "计算机组成原理期末高频必考模拟试卷C卷_2025真题风格补充_答案解析.docx"


def font_path(name):
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


FONT = ImageFont.truetype(font_path("arial.ttf") or font_path("simhei.ttf"), 24)
FONT_SMALL = ImageFont.truetype(font_path("arial.ttf") or font_path("simhei.ttf"), 18)
FONT_CN = ImageFont.truetype(font_path("simhei.ttf") or font_path("msyh.ttc") or font_path("arial.ttf"), 22)


def add_run(paragraph, text, bold=False, italic=False, size=11, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    add_run(p, text, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="9AA5B1", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_width(cell, inches):
    width = int(inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header_p, "计算机组成原理 高频必考模拟", size=9, color="666666")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, bold=True, size=20, color="0B2545")
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, subtitle, size=10.5, color="555555")
    p.paragraph_format.space_after = Pt(12)

    info = doc.add_table(rows=1, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    labels = ["考试时间：110 分钟", "满分：100 分", "姓名：__________", "学号：__________"]
    for i, text in enumerate(labels):
        set_cell_text(info.cell(0, i), text, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_fixed_width(info.cell(0, i), [1.55, 1.35, 1.7, 1.7][i])
    set_table_borders(info, "D6DAE0", "4")
    return doc


def add_section(doc, title):
    doc.add_paragraph(title, style="Heading 1")


def add_question(doc, number, text, points=None):
    p = doc.add_paragraph()
    prefix = f"{number}. "
    if points:
        prefix += f"（{points} 分）"
    add_run(p, prefix, bold=True)
    add_run(p, text)
    return p


def add_options(doc, options):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    add_run(p, "    ".join(options), size=10.5)


def add_blank_lines(doc, n=2):
    for _ in range(n):
        p = doc.add_paragraph()
        add_run(p, " " * 80)


def add_code(doc, code):
    for line in dedent(code).strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)


def draw_logic_hazard():
    img = Image.new("RGB", (1200, 430), "white")
    d = ImageDraw.Draw(img)
    blue = (32, 81, 120)
    gray = (110, 120, 130)
    black = (20, 20, 20)

    for y, label in [(85, "A"), (195, "B"), (305, "C")]:
        d.text((35, y - 15), label, font=FONT, fill=black)
        d.line((85, y, 260, y), fill=gray, width=4)
    d.rectangle((260, 55, 350, 115), outline=blue, width=4)
    d.text((285, 72), "NOT", font=FONT_SMALL, fill=black)
    d.line((350, 85, 460, 85), fill=gray, width=4)
    d.text((390, 52), "A'", font=FONT_SMALL, fill=black)
    d.line((175, 195, 460, 195), fill=gray, width=4)
    d.rectangle((460, 55, 590, 225), outline=blue, width=4)
    d.text((505, 122), "AND", font=FONT, fill=black)
    d.line((590, 140, 720, 140), fill=gray, width=4)
    d.line((175, 85, 460, 85), fill=gray, width=4)
    d.line((175, 305, 460, 305), fill=gray, width=4)
    d.rectangle((460, 250, 590, 365), outline=blue, width=4)
    d.text((505, 292), "AND", font=FONT, fill=black)
    d.line((590, 305, 720, 305), fill=gray, width=4)
    d.rectangle((720, 105, 865, 340), outline=blue, width=4)
    d.text((772, 210), "OR", font=FONT, fill=black)
    d.line((865, 222, 1110, 222), fill=gray, width=4)
    d.text((1120, 207), "F", font=FONT, fill=black)
    d.text((500, 25), "A'B", font=FONT_SMALL, fill=black)
    d.text((500, 370), "AC", font=FONT_SMALL, fill=black)
    path = FIG_DIR / "图1_组合逻辑险象电路.png"
    img.save(path)
    return path


def draw_kmap(path, answer=False):
    img = Image.new("RGB", (760, 560), "white")
    d = ImageDraw.Draw(img)
    x0, y0 = 120, 110
    cw, ch = 130, 85
    rows = ["00", "01", "11", "10"]
    cols = ["00", "01", "11", "10"]
    vals = [
        ["1", "X", "0", "1"],
        ["0", "1", "1", "0"],
        ["0", "1", "1", "0"],
        ["1", "X", "0", "1"],
    ]
    d.text((x0 + 35, 35), "CD", font=FONT_CN, fill=(0, 0, 0))
    d.text((25, y0 + 25), "AB", font=FONT_CN, fill=(0, 0, 0))
    for i, c in enumerate(cols):
        d.text((x0 + i * cw + 48, y0 - 40), c, font=FONT, fill=(0, 0, 0))
    for i, r in enumerate(rows):
        d.text((x0 - 55, y0 + i * ch + 28), r, font=FONT, fill=(0, 0, 0))
    for r in range(4):
        for c in range(4):
            left = x0 + c * cw
            top = y0 + r * ch
            d.rectangle((left, top, left + cw, top + ch), outline=(95, 110, 125), width=3)
            d.text((left + 56, top + 27), vals[r][c], font=FONT, fill=(0, 0, 0))
    if answer:
        red = (204, 61, 47)
        green = (32, 140, 95)
        d.rounded_rectangle((x0 + 0 * cw + 10, y0 + 0 * ch + 10, x0 + 3 * cw + cw - 10, y0 + 3 * ch + ch - 10), radius=28, outline=(210, 210, 210), width=1)
        d.rounded_rectangle((x0 + 0 * cw + 18, y0 + 0 * ch + 18, x0 + 3 * cw + cw - 18, y0 + 3 * ch + ch - 18), radius=22, outline=(210, 210, 210), width=1)
        # B'D' group wraps rows 00/10 and columns 00/10, shown as four corner marks.
        for r, c in [(0, 0), (0, 3), (3, 0), (3, 3)]:
            left = x0 + c * cw + 16
            top = y0 + r * ch + 16
            d.ellipse((left, top, left + 98, top + 55), outline=red, width=5)
        # BD group.
        d.rounded_rectangle((x0 + 1 * cw + 12, y0 + 1 * ch + 12, x0 + 2 * cw + cw - 12, y0 + 2 * ch + ch - 12), radius=24, outline=green, width=5)
        d.text((x0 + 20, y0 + 4 * ch + 24), "红圈: B'D'    绿圈: BD", font=FONT_CN, fill=(0, 0, 0))
    img.save(path)
    return path


def wave_points(levels, x0, y_low, y_high, step):
    pts = []
    for i, level in enumerate(levels):
        y = y_high if level else y_low
        x = x0 + i * step
        if i:
            pts.append((x, pts[-1][1]))
        pts.append((x, y))
        pts.append((x + step, y))
    return pts


def draw_waveform(path, answer=False):
    img = Image.new("RGB", (1100, 520), "white")
    d = ImageDraw.Draw(img)
    x0, step = 160, 80
    black = (20, 20, 20)
    gray = (140, 145, 150)
    blue = (32, 81, 120)
    clk = [0, 1] * 6
    din = [0, 1, 1, 0, 1, 0, 0, 1, 1, 0, 1, 1]
    qout = [0, 0, 1, 1, 0, 0, 1, 1, 0, 0, 1, 1]
    for i in range(13):
        x = x0 + i * step
        d.line((x, 70, x, 455), fill=(232, 235, 238), width=1)
        if i % 2 == 1:
            d.line((x, 70, x, 455), fill=(210, 217, 225), width=2)
            d.polygon([(x - 8, 68), (x + 8, 68), (x, 52)], fill=blue)
    for label, y in [("CLK", 125), ("D", 260), ("Q", 395)]:
        d.text((45, y - 18), label, font=FONT, fill=black)
        d.line((x0 - 20, y + 45, x0 + 12 * step + 20, y + 45), fill=(230, 230, 230), width=1)
    d.line(wave_points(clk, x0, 155, 95, step), fill=black, width=4)
    d.line(wave_points(din, x0, 290, 230, step), fill=black, width=4)
    if answer:
        d.line(wave_points(qout, x0, 425, 365, step), fill=(204, 61, 47), width=4)
    else:
        d.text((x0 + 60, 375), "请在此行画出 Q（初值 Q=0，上升沿触发）", font=FONT_CN, fill=(110, 110, 110))
    img.save(path)
    return path


def draw_state_101(path):
    img = Image.new("RGB", (950, 520), "white")
    d = ImageDraw.Draw(img)
    node_fill = (238, 244, 250)
    line = (38, 86, 132)
    nodes = {"S0/0": (130, 245), "S1/0": (350, 245), "S2/0": (570, 245), "S3/1": (790, 245)}
    for text, (x, y) in nodes.items():
        d.ellipse((x - 55, y - 40, x + 55, y + 40), fill=node_fill, outline=line, width=4)
        d.text((x - 33, y - 14), text, font=FONT_SMALL, fill=(0, 0, 0))

    def arrowhead(x1, y1, x2, y2):
        import math
        ang = math.atan2(y2 - y1, x2 - x1)
        size = 15
        pts = [
            (x2, y2),
            (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45)),
            (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45)),
        ]
        d.polygon(pts, fill=line)

    def arrow(a, b, label, offset=0):
        ax, ay = nodes[a]
        bx, by = nodes[b]
        if offset == 0:
            start = (ax + 55, ay)
            end = (bx - 55, by)
            d.line((*start, *end), fill=line, width=4)
            arrowhead(*start, *end)
            mx, my = (start[0] + end[0]) / 2, ay - 34
        else:
            left, right = min(ax, bx), max(ax, bx)
            top = ay + offset - 95
            bottom = ay + offset + 95
            if offset < 0:
                d.arc((left, top, right, bottom), 200, 340, fill=line, width=4)
                end = (bx - 45 if bx > ax else bx + 45, by - 15)
                arrowhead((ax + bx) / 2, ay + offset, *end)
                my = ay + offset - 35
            else:
                d.arc((left, top, right, bottom), 20, 160, fill=line, width=4)
                end = (bx - 45 if bx > ax else bx + 45, by + 15)
                arrowhead((ax + bx) / 2, ay + offset, *end)
                my = ay + offset + 12
            mx = (ax + bx) / 2
        d.text((mx - 12, my - 32), label, font=FONT_SMALL, fill=(0, 0, 0))

    arrow("S0/0", "S1/0", "1")
    arrow("S1/0", "S2/0", "0")
    arrow("S2/0", "S3/1", "1")
    arrow("S3/1", "S1/0", "1", offset=-155)
    arrow("S3/1", "S2/0", "0", offset=145)
    d.arc((65, 170, 195, 320), 70, 315, fill=line, width=4)
    arrowhead(86, 222, 74, 245)
    d.text((68, 150), "0", font=FONT_SMALL, fill=(0, 0, 0))
    d.arc((295, 170, 405, 320), 210, 35, fill=line, width=4)
    arrowhead(386, 219, 398, 245)
    d.text((345, 148), "1", font=FONT_SMALL, fill=(0, 0, 0))
    d.text((80, 455), "Moore 型 101 序列检测器：状态名/输出，允许重叠", font=FONT_CN, fill=(0, 0, 0))
    img.save(path)
    return path


def draw_majority_answer(path):
    img = Image.new("RGB", (1080, 420), "white")
    d = ImageDraw.Draw(img)
    blue = (32, 81, 120)
    gray = (115, 125, 135)
    black = (20, 20, 20)
    for y, label in [(70, "A"), (190, "B"), (310, "C")]:
        d.text((35, y - 15), label, font=FONT, fill=black)
        d.line((85, y, 280, y), fill=gray, width=4)
    labels = [("AB", 360, 55), ("AC", 360, 165), ("BC", 360, 275)]
    connections = [
        ((120, 70, 360, 85), (120, 190, 360, 115)),
        ((120, 70, 360, 195), (120, 310, 360, 225)),
        ((120, 190, 360, 305), (120, 310, 360, 335)),
    ]
    for (name, x, y), con in zip(labels, connections):
        d.line(con[0], fill=gray, width=3)
        d.line(con[1], fill=gray, width=3)
        d.rectangle((x, y, x + 125, y + 70), outline=blue, width=4)
        d.text((x + 38, y + 22), "AND", font=FONT_SMALL, fill=black)
        d.line((x + 125, y + 35, 650, y + 35), fill=gray, width=4)
        d.text((x + 142, y + 8), name, font=FONT_SMALL, fill=black)
    d.rectangle((650, 95, 800, 325), outline=blue, width=4)
    d.text((704, 200), "OR", font=FONT, fill=black)
    d.line((800, 210, 1010, 210), fill=gray, width=4)
    d.text((1020, 195), "Y", font=FONT, fill=black)
    img.save(path)
    return path


def draw_sync_counter_circuit(path):
    img = Image.new("RGB", (1800, 920), "white")
    d = ImageDraw.Draw(img)
    blue = (28, 78, 118)
    gray = (82, 96, 110)
    light = (232, 237, 242)
    black = (20, 20, 20)
    title_font = ImageFont.truetype(font_path("simhei.ttf") or font_path("msyh.ttc"), 36)
    label_font = ImageFont.truetype(font_path("arial.ttf") or font_path("simhei.ttf"), 30)
    small_font = ImageFont.truetype(font_path("arial.ttf") or font_path("simhei.ttf"), 24)
    cn_small = ImageFont.truetype(font_path("simhei.ttf") or font_path("msyh.ttc"), 25)

    def line(points, width=5, fill=gray):
        d.line(points, fill=fill, width=width, joint="curve")

    def dot(x, y, r=7):
        d.ellipse((x - r, y - r, x + r, y + r), fill=black)

    def ff(x, y, name):
        w, h = 220, 210
        d.rectangle((x, y, x + w, y + h), outline=blue, width=6, fill="white")
        d.text((x + 78, y + 78), "JK", font=label_font, fill=black)
        d.text((x + 15, y + 92), "J=K", font=small_font, fill=black)
        # Clock marker.
        tri = [(x, y + h - 72), (x, y + h - 32), (x + 28, y + h - 52)]
        d.polygon(tri, outline=blue, fill="white")
        d.text((x + 86, y + h - 36), name, font=small_font, fill=black)
        # Output pin.
        line((x + w, y + h // 2, x + w + 55, y + h // 2), width=5)
        d.text((x + w + 66, y + h // 2 - 18), name, font=small_font, fill=black)
        return {
            "T": (x, y + 105),
            "CLK": (x, y + h - 52),
            "Q": (x + w + 55, y + h // 2),
            "Q_SRC": (x + w, y + h // 2),
            "x": x,
            "y": y,
            "w": w,
            "h": h,
        }

    def and_gate(x, y, label, inputs=2):
        h = 96 if inputs == 3 else 82
        d.rectangle((x, y, x + 104, y + h), outline=blue, width=5, fill="white")
        d.text((x + 38, y + h // 2 - 18), "&", font=label_font, fill=black)
        d.text((x + 2, y + 90), label, font=cn_small, fill=black)
        if inputs == 3:
            return {
                "in1": (x, y + 20),
                "in2": (x, y + 48),
                "in3": (x, y + 76),
                "out": (x + 104, y + 48),
            }
        return {
            "in1": (x, y + 22),
            "in2": (x, y + 60),
            "out": (x + 104, y + 41),
        }

    d.text((360, 40), "三位同步二进制计数器分析电路：所有触发器同一 CP 上升沿触发", font=title_font, fill=black)

    x_bus_y = 130
    cp_y = 790
    left_bus = 90
    right_bus = 1655
    d.text((35, x_bus_y - 20), "x", font=label_font, fill=black)
    line((left_bus, x_bus_y, right_bus, x_bus_y), width=6)
    d.text((35, cp_y - 20), "CP", font=label_font, fill=black)
    line((left_bus, cp_y, right_bus - 80, cp_y), width=6)

    ff0 = ff(250, 270, "Q0")
    ff1 = ff(720, 270, "Q1")
    ff2 = ff(1190, 270, "Q2")
    gate1 = and_gate(555, 540, "T1=xQ0", inputs=2)
    gate2 = and_gate(1018, 522, "T2=xQ0Q1", inputs=3)

    # CP bus to every FF clock input.
    for f in (ff0, ff1, ff2):
        cx, cy = f["CLK"]
        line((cx, cp_y, cx, cy), width=5)
        dot(cx, cp_y)

    # T0 = x feeds J0=K0.
    t0_x = ff0["T"][0] - 72
    line((t0_x, x_bus_y, t0_x, ff0["T"][1]), width=5)
    line((t0_x, ff0["T"][1], *ff0["T"]), width=5)
    dot(t0_x, x_bus_y)
    d.text((ff0["x"] - 70, ff0["y"] - 42), "T0=J0=K0=x", font=small_font, fill=black)

    # x and Q0 to first AND.
    x_drop1 = gate1["in1"][0] + 22
    line((x_drop1, x_bus_y, x_drop1, gate1["in1"][1]), width=4)
    line((x_drop1, gate1["in1"][1], *gate1["in1"]), width=4)
    dot(x_drop1, x_bus_y)
    line((ff0["Q"][0], ff0["Q"][1], ff0["Q"][0] + 70, ff0["Q"][1]), width=5)
    q0_branch_x = ff0["Q"][0] + 60
    line((q0_branch_x, ff0["Q"][1], q0_branch_x, gate1["in2"][1]), width=4)
    line((q0_branch_x, gate1["in2"][1], *gate1["in2"]), width=4)
    dot(q0_branch_x, ff0["Q"][1])

    # Gate1 output to T1 = J1 = K1.
    g1_out = gate1["out"]
    feed1_x = ff1["T"][0] - 58
    line((*g1_out, feed1_x, g1_out[1]), width=5)
    line((feed1_x, g1_out[1], feed1_x, ff1["T"][1]), width=5)
    line((feed1_x, ff1["T"][1], *ff1["T"]), width=5)
    d.text((ff1["x"] - 84, ff1["y"] - 42), "T1=J1=K1=xQ0", font=small_font, fill=black)

    # x, Q0, Q1 to second AND.
    x_drop2 = gate2["in1"][0] + 18
    line((x_drop2, x_bus_y, x_drop2, gate2["in1"][1]), width=4)
    line((x_drop2, gate2["in1"][1], *gate2["in1"]), width=4)
    dot(x_drop2, x_bus_y)
    q1_branch_x = ff1["Q"][0] + 60
    line((ff1["Q"][0], ff1["Q"][1], ff1["Q"][0] + 70, ff1["Q"][1]), width=5)
    line((q1_branch_x, ff1["Q"][1], q1_branch_x, gate2["in3"][1]), width=4)
    line((q1_branch_x, gate2["in3"][1], *gate2["in3"]), width=4)
    dot(q1_branch_x, ff1["Q"][1])
    q0_to_g2_y = gate2["in2"][1]
    line((q0_branch_x, ff0["Q"][1] + 70, q0_branch_x, q0_to_g2_y), width=3)
    line((q0_branch_x, q0_to_g2_y, gate2["in2"][0], q0_to_g2_y), width=3)

    # Gate2 output to T2 = J2 = K2.
    g2_out = gate2["out"]
    feed2_x = ff2["T"][0] - 58
    line((*g2_out, feed2_x, g2_out[1]), width=5)
    line((feed2_x, g2_out[1], feed2_x, ff2["T"][1]), width=5)
    line((feed2_x, ff2["T"][1], *ff2["T"]), width=5)
    d.text((ff2["x"] - 112, ff2["y"] - 42), "T2=J2=K2=xQ0Q1", font=small_font, fill=black)

    # Output.
    line((ff2["Q"][0], ff2["Q"][1], 1615, ff2["Q"][1]), width=5)
    d.text((1630, ff2["Q"][1] - 18), "Z=Q2Q1Q0", font=small_font, fill=black)

    # Subtle input pin labels for the gates.
    d.text((gate1["in1"][0] - 50, gate1["in1"][1] - 18), "x", font=small_font, fill=black)
    d.text((gate1["in2"][0] - 72, gate1["in2"][1] - 18), "Q0", font=small_font, fill=black)
    d.text((gate2["in1"][0] - 50, gate2["in1"][1] - 18), "x", font=small_font, fill=black)
    d.text((gate2["in2"][0] - 72, gate2["in2"][1] - 18), "Q0", font=small_font, fill=black)
    d.text((gate2["in3"][0] - 72, gate2["in3"][1] - 18), "Q1", font=small_font, fill=black)

    # Downsample for anti-aliased clean lines in Word/PDF.
    img = img.resize((1200, 613), Image.Resampling.LANCZOS)
    img.save(path)
    return path


def draw_seq1001_state(path):
    img = Image.new("RGB", (1100, 560), "white")
    d = ImageDraw.Draw(img)
    blue = (32, 81, 120)
    fill = (238, 244, 250)
    black = (20, 20, 20)
    nodes = {"A": (140, 270), "B": (380, 270), "C": (620, 270), "D": (860, 270)}
    desc = {
        "A": "无匹配",
        "B": "已收 1",
        "C": "已收 10",
        "D": "已收 100",
    }
    for s, (x, y) in nodes.items():
        d.ellipse((x - 62, y - 42, x + 62, y + 42), outline=blue, fill=fill, width=4)
        d.text((x - 12, y - 20), s, font=FONT, fill=black)
        d.text((x - 40, y + 12), desc[s], font=FONT_SMALL, fill=black)

    def line_arrow(start, end, label, above=True):
        import math
        x1, y1 = start
        x2, y2 = end
        d.line((x1, y1, x2, y2), fill=blue, width=4)
        ang = math.atan2(y2 - y1, x2 - x1)
        size = 15
        d.polygon([
            (x2, y2),
            (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45)),
            (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45)),
        ], fill=blue)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2 + (-38 if above else 20)
        d.text((mx - 22, my), label, font=FONT_SMALL, fill=black)

    line_arrow((202, 270), (318, 270), "1/0")
    line_arrow((442, 270), (558, 270), "0/0")
    line_arrow((682, 270), (798, 270), "0/0")
    # D on 1 returns to B with output.
    d.arc((380, 120, 860, 420), 205, 335, fill=blue, width=4)
    d.text((604, 116), "1/1", font=FONT_SMALL, fill=black)
    d.polygon([(422, 228), (438, 224), (431, 242)], fill=blue)
    # Self/return arcs.
    d.arc((70, 185, 210, 350), 70, 315, fill=blue, width=4)
    d.text((78, 160), "0/0", font=FONT_SMALL, fill=black)
    d.arc((315, 175, 445, 345), 210, 35, fill=blue, width=4)
    d.text((375, 156), "1/0", font=FONT_SMALL, fill=black)
    d.arc((380, 355, 620, 500), 25, 160, fill=blue, width=4)
    d.text((490, 456), "1/0", font=FONT_SMALL, fill=black)
    d.line((860, 312, 140, 312), fill=blue, width=4)
    d.polygon([(140, 312), (158, 302), (158, 322)], fill=blue)
    d.text((500, 320), "0/0", font=FONT_SMALL, fill=black)
    d.text((150, 45), "Mealy 型可重叠 1001 序列检测器原始状态图（边标注 x/z）", font=FONT_CN, fill=black)
    img.save(path)
    return path


def draw_minimized_state(path):
    img = Image.new("RGB", (850, 420), "white")
    d = ImageDraw.Draw(img)
    blue = (32, 81, 120)
    fill = (238, 244, 250)
    black = (20, 20, 20)
    nodes = {"P": (150, 220), "Q": (405, 220), "R": (660, 220)}
    labels = {"P": "{A,B,D}", "Q": "{C,F}", "R": "{E,G}"}
    for s, (x, y) in nodes.items():
        d.ellipse((x - 70, y - 44, x + 70, y + 44), outline=blue, fill=fill, width=4)
        d.text((x - 10, y - 25), s, font=FONT, fill=black)
        d.text((x - 46, y + 10), labels[s], font=FONT_SMALL, fill=black)

    def arr(x1, y1, x2, y2, text, above=True):
        import math
        d.line((x1, y1, x2, y2), fill=blue, width=4)
        ang = math.atan2(y2 - y1, x2 - x1)
        size = 15
        d.polygon([
            (x2, y2),
            (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45)),
            (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45)),
        ], fill=blue)
        d.text(((x1 + x2) / 2 - 22, (y1 + y2) / 2 + (-34 if above else 16)), text, font=FONT_SMALL, fill=black)

    arr(220, 220, 335, 220, "1/0")
    arr(475, 220, 590, 220, "1/0")
    d.arc((80, 145, 220, 295), 70, 315, fill=blue, width=4)
    d.text((75, 122), "0/0", font=FONT_SMALL, fill=black)
    d.arc((335, 145, 475, 295), 70, 315, fill=blue, width=4)
    d.text((365, 122), "0/0", font=FONT_SMALL, fill=black)
    d.line((660, 264, 150, 264), fill=blue, width=4)
    d.polygon([(150, 264), (168, 254), (168, 274)], fill=blue)
    d.text((378, 276), "0/0, 1/1", font=FONT_SMALL, fill=black)
    d.text((150, 40), "化简后状态图：P={A,B,D}, Q={C,F}, R={E,G}", font=FONT_CN, fill=black)
    img.save(path)
    return path


def make_figures():
    return {
        "logic": draw_logic_hazard(),
        "kmap_q": draw_kmap(FIG_DIR / "图2_卡诺图题目.png", False),
        "kmap_a": draw_kmap(FIG_DIR / "图2_卡诺图答案.png", True),
        "wave_q": draw_waveform(FIG_DIR / "图3_D触发器波形题目.png", False),
        "wave_a": draw_waveform(FIG_DIR / "图3_D触发器波形答案.png", True),
        "state": draw_state_101(FIG_DIR / "图4_101序列检测器状态图答案.png"),
        "majority": draw_majority_answer(FIG_DIR / "图5_三人表决器答案电路.png"),
        "sync_counter": draw_sync_counter_circuit(FIG_DIR / "图6_JK同步计数器分析电路.png"),
        "seq1001": draw_seq1001_state(FIG_DIR / "图7_1001序列检测器原始状态图.png"),
        "min_state": draw_minimized_state(FIG_DIR / "图8_状态表化简结果图.png"),
    }


CHOICES = [
    ("下列关于数字信号的说法正确的是（  ）。", ["A. 数值连续且时间连续", "B. 通常在时间和数值上离散", "C. 不能进行逻辑运算", "D. 抗干扰能力一定弱于模拟信号"], "B"),
    ("组合逻辑电路的输出取决于（  ）。", ["A. 当前输入", "B. 当前输入和历史状态", "C. 时钟边沿", "D. 触发器现态"], "A"),
    ("8421 BCD 码中，非法码的个数是（  ）。", ["A. 4", "B. 6", "C. 8", "D. 10"], "B"),
    ("二进制数转换为格雷码时，格雷码最高位（  ）。", ["A. 恒为 0", "B. 恒为 1", "C. 等于二进制最高位", "D. 等于二进制最低位"], "C"),
    ("与非门构成的基本 RS 触发器通常属于（  ）。", ["A. 高电平有效", "B. 低电平有效", "C. 边沿有效", "D. 不允许保持"], "B"),
    ("JK 触发器在 J=K=1 时，次态为（  ）。", ["A. 保持", "B. 置 0", "C. 置 1", "D. 翻转"], "D"),
    ("卡诺图化简时，正确的圈组原则是（  ）。", ["A. 圈越小越好", "B. 每个圈中格数必须为 3 的倍数", "C. 圈组尽量大且数量尽量少", "D. 无关项必须全部圈入"], "C"),
    ("Moore 型状态机的输出（  ）。", ["A. 只与现态有关", "B. 只与输入有关", "C. 与现态和输入均无关", "D. 一定不能用 VHDL 描述"], "A"),
    ("VHDL 中 signal 赋值通常使用（  ）。", ["A. :=", "B. <=", "C. ==", "D. =>"], "B"),
    ("描述时序逻辑时，常用的上升沿判断语句是（  ）。", ["A. if clk='1' then", "B. if rising_edge(clk) then", "C. wait clk", "D. case clk is"], "B"),
]


FILL = [
    "R 进制数转换为十进制数常用的方法是________。",
    "8 位补码整数的表示范围是________到________。",
    "n 个变量共有________个最小项。",
    "卡诺图中的无关项 X 在化简时________（必须圈入/可用可不用）。",
    "三态门的三个输出状态分别是 0、1 和________。",
    "OC 门输出端通常需要外接________电阻，并可实现________功能。",
    "静态 1 险象是指理论输出应保持 1，实际输出短暂变为________。",
    "若状态数为 M，触发器个数 n 应满足________。",
    "VHDL 程序基本结构通常包括 entity 和________。",
    "VHDL case 语句中建议写________分支以覆盖未列状态。",
]


def build_question_doc(figs):
    doc = setup_doc("计算机组成原理期末高频必考模拟试卷 A 卷", "依据《计算机组成原理期末复习资料》高频题型与必考框架编制")

    add_section(doc, "一、选择题（每题 2 分，共 20 分）")
    for i, (stem, opts, _) in enumerate(CHOICES, 1):
        add_question(doc, i, stem)
        add_options(doc, opts)

    add_section(doc, "二、填空题（每空 1 分，共 12 分）")
    for i, stem in enumerate(FILL, 1):
        add_question(doc, i, stem)

    add_section(doc, "三、简答题（每题 5 分，共 20 分）")
    short = [
        "简述组合逻辑电路与时序逻辑电路的区别，并各举一个典型例子。",
        "简述三态门和 OC 门的用途及使用注意事项。",
        "说明用卡诺图化简含无关项逻辑函数的一般步骤。",
        "比较 VHDL 中 signal 与 variable、if 与 case 的主要区别。",
    ]
    for i, stem in enumerate(short, 1):
        add_question(doc, i, stem, 5)
        add_blank_lines(doc, 2)

    add_section(doc, "四、计算与分析题（共 26 分）")
    add_question(doc, 1, "完成下列数制与编码计算：① 将 (110101.101)2 转换为十进制；② 将 -45 写成 8 位原码、反码、补码；③ 将十进制 93 写成 8421 BCD 码和余三码；④ 将二进制 101101 转换为格雷码。", 8)
    add_blank_lines(doc, 3)
    add_question(doc, 2, "用图 2 所示 4 变量卡诺图化简 F(A,B,C,D)，其中 X 表示无关项，写出最简与或式。", 7)
    doc.add_picture(str(figs["kmap_q"]), width=Inches(4.8))
    add_blank_lines(doc, 2)
    add_question(doc, 3, "图 1 所示组合电路对应 F=A'B+AC。判断该电路是否存在静态 1 险象；若存在，指出发生条件并给出消除后的表达式。", 5)
    doc.add_picture(str(figs["logic"]), width=Inches(5.7))
    add_blank_lines(doc, 2)
    add_question(doc, 4, "上升沿 D 触发器初值 Q=0，输入波形如图 3。请画出 Q 的波形，并说明采样规律。", 6)
    doc.add_picture(str(figs["wave_q"]), width=Inches(5.8))
    add_blank_lines(doc, 2)

    add_section(doc, "五、设计与 VHDL 题（共 22 分）")
    add_question(doc, 1, "设计三人多数表决电路。要求：设输入为 A、B、C，输出为 Y；列真值表，写出最简表达式，并画出逻辑实现或写出 VHDL。", 10)
    add_blank_lines(doc, 4)
    add_question(doc, 2, "设计一个 Moore 型“101”序列检测器，允许重叠。要求：给出状态含义、状态转移关系，并写出 VHDL 三段式状态机核心代码。", 12)
    add_blank_lines(doc, 5)

    doc.save(QUESTION_DOCX)
    return QUESTION_DOCX


def add_answer_item(doc, number, title, body):
    p = doc.add_paragraph()
    add_run(p, f"{number}. {title}", bold=True, size=11.5, color="0B2545")
    for para in dedent(body).strip().split("\n"):
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            add_run(p, para.strip(), size=10.5)


def build_answer_doc(figs):
    doc = setup_doc("计算机组成原理期末高频必考模拟试卷 A 卷 答案解析", "答案按采分点编写，便于自测后核对")

    add_section(doc, "一、选择题答案")
    table = doc.add_table(rows=2, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i in range(10):
        set_cell_text(table.cell(0, i), str(i + 1), bold=True)
        set_cell_text(table.cell(1, i), CHOICES[i][2], bold=True)
        set_fixed_width(table.cell(0, i), 0.55)
        set_fixed_width(table.cell(1, i), 0.55)
        shade_cell(table.cell(0, i), "E8EEF5")
    set_table_borders(table)

    add_section(doc, "二、填空题答案")
    fill_answers = [
        "按权展开",
        "-128；+127",
        "2^n",
        "可用可不用",
        "高阻态 Z",
        "上拉；线与",
        "0",
        "2^n >= M",
        "architecture",
        "when others",
    ]
    for i, ans in enumerate(fill_answers, 1):
        add_answer_item(doc, i, "", ans)

    add_section(doc, "三、简答题要点")
    add_answer_item(doc, 1, "组合逻辑与时序逻辑", """
    组合逻辑：输出只由当前输入决定，无记忆元件，如译码器、数据选择器、全加器。
    时序逻辑：输出与当前输入及现态有关，含触发器等存储元件，如寄存器、计数器、序列检测器。
    """)
    add_answer_item(doc, 2, "三态门与 OC 门", """
    三态门输出有 0、1、高阻态 Z，常用于总线分时传送；同一总线任一时刻只能有一个有效输出，否则可能冲突。
    OC 门为集电极开路门，输出端需外接上拉电阻，可实现线与功能。
    """)
    add_answer_item(doc, 3, "卡诺图含无关项化简", """
    先按最小项填 1，按约束项填 X；圈组大小取 1、2、4、8 等 2 的幂；圈组尽量大、数量尽量少；X 有利于扩大圈组时可当作 1，否则当作 0；最后只保留圈内不变变量。
    """)
    add_answer_item(doc, 4, "VHDL 区别", """
    signal 用 <= 赋值，通常在进程挂起后更新；variable 用 := 赋值，在进程内部立即更新。
    if 语句有优先级，适合条件优先判断；case 分支并列，适合译码器和状态机，并建议写 when others。
    """)

    add_section(doc, "四、计算与分析题答案")
    add_answer_item(doc, 1, "数制与编码", """
    ① (110101.101)2 = 32+16+4+1+0.5+0.125 = 53.625。
    ② +45 的 8 位二进制为 00101101，因此 -45 原码=10101101，反码=11010010，补码=11010011。
    ③ 93 的 8421 BCD=1001 0011；余三码为 9->1100，3->0110，即 1100 0110。
    ④ 101101 转格雷码：最高位不变，其余位为相邻二进制位异或，得 111011。
    """)
    add_answer_item(doc, 2, "卡诺图化简", """
    图中 1 项为 m(0,2,5,7,8,10,13,15)，无关项为 d(1,9)。
    四角 1 可圈成 B'D'；中间 2×2 的 1 可圈成 BD。
    所以 F = B'D' + BD，即 B 与 D 的同或。
    """)
    doc.add_picture(str(figs["kmap_a"]), width=Inches(4.8))
    add_answer_item(doc, 3, "静态 1 险象", """
    F=A'B+AC。当 B=1、C=1 且 A 从 0 到 1 或从 1 到 0 变化时，理论上 F 应保持 1，但两个乘积项的传输延迟不同，可能出现短暂 0。
    因此存在静态 1 险象。增加覆盖相邻 1 的冗余项 BC，消除后 F=A'B+AC+BC。
    """)
    doc.add_picture(str(figs["logic"]), width=Inches(5.7))
    add_answer_item(doc, 4, "D 触发器波形", """
    上升沿 D 触发器只在 CLK 上升沿采样 D，并把采样值送到 Q；两个上升沿之间 Q 保持不变。
    本题初值 Q=0，按图中各上升沿采样得到的 Q 波形如下。
    """)
    doc.add_picture(str(figs["wave_a"]), width=Inches(5.8))

    add_section(doc, "五、设计与 VHDL 题答案")
    add_answer_item(doc, 1, "三人多数表决电路", """
    真值规律：A、B、C 中至少两个为 1 时 Y=1，否则 Y=0。
    最简表达式：Y = AB + AC + BC。
    """)
    doc.add_picture(str(figs["majority"]), width=Inches(5.7))
    add_code(doc, """
    library IEEE;
    use IEEE.STD_LOGIC_1164.ALL;

    entity majority3 is
      port(
        A, B, C : in  std_logic;
        Y       : out std_logic
      );
    end majority3;

    architecture behavior of majority3 is
    begin
      Y <= (A and B) or (A and C) or (B and C);
    end behavior;
    """)

    add_answer_item(doc, 2, "Moore 型 101 序列检测器", """
    状态含义：S0 表示未匹配；S1 表示已匹配到 1；S2 表示已匹配到 10；S3 表示已匹配到 101，输出 y=1。
    允许重叠时，S3 遇到输入 1 转到 S1，遇到输入 0 转到 S2。
    """)
    doc.add_picture(str(figs["state"]), width=Inches(5.9))
    add_code(doc, """
    library IEEE;
    use IEEE.STD_LOGIC_1164.ALL;

    entity seq101_moore is
      port(
        clk, rst : in  std_logic;
        x        : in  std_logic;
        y        : out std_logic
      );
    end seq101_moore;

    architecture behavior of seq101_moore is
      type state_type is (S0, S1, S2, S3);
      signal state, next_state : state_type;
    begin
      process(clk, rst)
      begin
        if rst = '1' then
          state <= S0;
        elsif rising_edge(clk) then
          state <= next_state;
        end if;
      end process;

      process(state, x)
      begin
        case state is
          when S0 =>
            if x = '1' then next_state <= S1;
            else next_state <= S0;
            end if;
          when S1 =>
            if x = '0' then next_state <= S2;
            else next_state <= S1;
            end if;
          when S2 =>
            if x = '1' then next_state <= S3;
            else next_state <= S0;
            end if;
          when S3 =>
            if x = '1' then next_state <= S1;
            else next_state <= S2;
            end if;
          when others =>
            next_state <= S0;
        end case;
      end process;

      process(state)
      begin
        case state is
          when S3 => y <= '1';
          when others => y <= '0';
        end case;
      end process;
    end behavior;
    """)

    doc.save(ANSWER_DOCX)
    return ANSWER_DOCX


def add_state_table(doc, rows, title=None):
    if title:
        p = doc.add_paragraph()
        add_run(p, title, bold=True, size=10.5)
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["现态", "x=0 次态/输出", "x=1 次态/输出"]
    widths = [1.2, 2.25, 2.25]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "E8EEF5")
        set_fixed_width(table.cell(0, i), widths[i])
    for r, row in enumerate(rows, 1):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text)
            set_fixed_width(table.cell(r, c), widths[c])
    set_table_borders(table)
    return table


def build_sync_doc(figs):
    doc = setup_doc("计算机组成原理期末高频必考模拟试卷 B 卷", "提高补充卷：同步时序大题、触发器转换、MSI 实现与 VHDL")
    add_section(doc, "一、同步时序电路分析（18 分）")
    add_question(doc, 1, "分析图 1 所示同步时序逻辑电路，写出各 JK 触发器的激励函数、次态方程、输出函数，列出 x=1 时的状态转换序列，并说明该电路的功能。设初态 Q2Q1Q0=000。", 18)
    doc.add_picture(str(figs["sync_counter"]), width=Inches(6.2))
    add_blank_lines(doc, 4)

    add_section(doc, "二、可重叠序列检测器原始状态图与状态表（18 分）")
    add_question(doc, 1, "试画出“1001”可重叠序列检测器的原始状态图和原始状态表。输入为 x，输出为 z；当且仅当当前输入使最近 4 位形成 1001 时，z=1，其余时刻 z=0。建议采用 Mealy 型设计。", 18)
    add_blank_lines(doc, 5)

    add_section(doc, "三、原始状态表化简（18 分）")
    add_question(doc, 1, "化简下列 Mealy 型原始状态表，写出等价状态集合、最小化状态表，并画出化简后的状态图。", 18)
    rows = [
        ("A", "B/0", "C/0"),
        ("B", "A/0", "F/0"),
        ("C", "F/0", "G/0"),
        ("D", "A/0", "C/0"),
        ("E", "A/0", "A/1"),
        ("F", "C/0", "E/0"),
        ("G", "A/0", "B/1"),
    ]
    add_state_table(doc, rows)
    add_blank_lines(doc, 4)

    add_section(doc, "四、触发器转换与指定门实现（16 分）")
    add_question(doc, 1, "用 JK 触发器实现 D 触发器功能。要求：列出转换表，求出 J、K 关于 D 和 Qn 的最简表达式。", 8)
    add_blank_lines(doc, 3)
    add_question(doc, 2, "将逻辑函数 F(A,B,C)=Σm(1,2,3,5,7) 化为最简与或式，并进一步写成与非-与非实现形式。", 8)
    add_blank_lines(doc, 3)

    add_section(doc, "五、MSI 组合逻辑实现（16 分）")
    add_question(doc, 1, "用 3 线-8 线译码器和必要的门电路实现 F(A,B,C)=Σm(0,2,5,7)。说明译码器输出端与最小项的对应关系，并写出连接方式。", 8)
    add_blank_lines(doc, 3)
    add_question(doc, 2, "用 8 选 1 数据选择器实现 F(A,B,C,D)=Σm(1,3,4,11,12,13,15)。取 A、B、C 为选择变量，列出各数据输入端 D0-D7 应接 0、1、D 或 D'。", 8)
    add_blank_lines(doc, 3)

    add_section(doc, "六、VHDL 补全与说明（14 分）")
    add_question(doc, 1, "根据第二题的“1001”可重叠序列检测器，补全下面 VHDL 状态机中的空缺，并说明 process 的作用。", 14)
    add_code(doc, """
    library IEEE;
    use IEEE.STD_LOGIC_1164.ALL;

    entity seq1001 is
      port(
        clk, rst : in  std_logic;
        x        : in  std_logic;
        z        : out std_logic
      );
    end seq1001;

    architecture behavior of seq1001 is
      type state_type is (A, B, C, D);
      signal state, next_state : state_type;
    begin
      process(clk, rst)
      begin
        if rst = '1' then
          state <= ________;
        elsif rising_edge(clk) then
          state <= ________;
        end if;
      end process;

      process(state, x)
      begin
        z <= '0';
        case state is
          when A =>
            if x = '1' then next_state <= ________;
            else next_state <= ________;
            end if;
          when B =>
            if x = '0' then next_state <= ________;
            else next_state <= ________;
            end if;
          when C =>
            if x = '0' then next_state <= ________;
            else next_state <= ________;
            end if;
          when D =>
            if x = '1' then
              next_state <= ________;
              z <= ________;
            else
              next_state <= ________;
            end if;
          when others =>
            next_state <= A;
        end case;
      end process;
    end behavior;
    """)
    doc.save(SYNC_DOCX)
    return SYNC_DOCX


def build_sync_answer_doc(figs):
    doc = setup_doc("计算机组成原理期末高频必考模拟试卷 B 卷 答案解析", "提高补充卷答案")

    add_section(doc, "一、同步时序电路分析")
    add_answer_item(doc, 1, "激励函数、次态方程与功能", """
    由图可读出：J0=K0=x，J1=K1=xQ0，J2=K2=xQ1Q0。
    JK 触发器特性方程为 Q(n+1)=JQ' + K'Q。当 J=K=T 时等价为 T 触发器，Q(n+1)=Q⊕T。
    因此 Q0+=Q0⊕x；Q1+=Q1⊕(xQ0)；Q2+=Q2⊕(xQ1Q0)。
    输出函数 Z=Q2Q1Q0，可作为计数到 111 的终端状态指示。
    当 x=0 时三个触发器均保持；当 x=1 时，状态序列为 000→001→010→011→100→101→110→111→000。
    所以该电路是带使能 x 的三位同步二进制加 1 计数器，模为 8。
    """)
    doc.add_picture(str(figs["sync_counter"]), width=Inches(6.2))

    add_section(doc, "二、1001 可重叠序列检测器")
    add_answer_item(doc, 1, "状态含义与转移", """
    采用 Mealy 型。A 表示未匹配；B 表示已匹配到 1；C 表示已匹配到 10；D 表示已匹配到 100。
    在 D 状态输入 1 时检测到 1001，输出 z=1。因为允许重叠，检测后最近一位仍为 1，所以下一状态转到 B。
    """)
    rows = [
        ("A", "A/0", "B/0"),
        ("B", "C/0", "B/0"),
        ("C", "D/0", "B/0"),
        ("D", "A/0", "B/1"),
    ]
    add_state_table(doc, rows, "原始状态表")
    doc.add_picture(str(figs["seq1001"]), width=Inches(6.2))

    add_section(doc, "三、状态表化简")
    add_answer_item(doc, 1, "等价状态集合", """
    先按输出对划分：A、B、C、D、F 的输出对均为 0/0；E、G 的输出对均为 0/1。
    继续按次态所在集合划分，可得 A、B、D 等价；C、F 等价；E、G 等价。
    因此最小状态集合为 P={A,B,D}，Q={C,F}，R={E,G}。
    """)
    rows = [
        ("P={A,B,D}", "P/0", "Q/0"),
        ("Q={C,F}", "Q/0", "R/0"),
        ("R={E,G}", "P/0", "P/1"),
    ]
    add_state_table(doc, rows, "最小化状态表")
    doc.add_picture(str(figs["min_state"]), width=Inches(5.6))

    add_section(doc, "四、触发器转换与指定门实现")
    add_answer_item(doc, 1, "JK 触发器实现 D 触发器", """
    目标 D 触发器满足 Q(n+1)=D。
    转换表：Qn=0,D=0 时 Q+=0，JK 可取 J=0,K=X；Qn=0,D=1 时 Q+=1，J=1,K=X；Qn=1,D=0 时 Q+=0，J=X,K=1；Qn=1,D=1 时 Q+=1，J=X,K=0。
    化简得 J=D，K=D'。因此把 D 接到 J，D 取反后接到 K，即可用 JK 触发器实现 D 触发器。
    """)
    add_answer_item(doc, 2, "与非-与非实现", """
    F(A,B,C)=Σm(1,2,3,5,7)。由卡诺图可化简为 F=C+AB'。
    与非-与非形式：F=((C)'·(AB')')'，也可写作 F=NAND(NAND(C,C), NAND(A,B'))，其中 B'=NAND(B,B)。
    """)

    add_section(doc, "五、MSI 组合逻辑实现")
    add_answer_item(doc, 1, "3-8 译码器实现", """
    若译码器输出 Y0-Y7 分别对应最小项 m0-m7，且为高有效，则 F=Y0+Y2+Y5+Y7。
    连接方式：A、B、C 接译码器地址输入，取 Y0、Y2、Y5、Y7 送入或门得到 F。
    若使用低有效译码器，则应按器件有效电平改用与非门汇总。
    """)
    add_answer_item(doc, 2, "8 选 1 数据选择器实现", """
    取 A、B、C 为选择变量，逐组观察 D=0 和 D=1 两种情况。
    ABC=000：m0 不取、m1 取，D0=D；001：m2 不取、m3 取，D1=D；010：m4 取、m5 不取，D2=D'；011：m6、m7 均不取，D3=0。
    ABC=100：m8、m9 均不取，D4=0；101：m10 不取、m11 取，D5=D；110：m12、m13 均取，D6=1；111：m14 不取、m15 取，D7=D。
    因此 D0=D，D1=D，D2=D'，D3=0，D4=0，D5=D，D6=1，D7=D。
    """)

    add_section(doc, "六、VHDL 补全")
    add_code(doc, """
    library IEEE;
    use IEEE.STD_LOGIC_1164.ALL;

    entity seq1001 is
      port(
        clk, rst : in  std_logic;
        x        : in  std_logic;
        z        : out std_logic
      );
    end seq1001;

    architecture behavior of seq1001 is
      type state_type is (A, B, C, D);
      signal state, next_state : state_type;
    begin
      process(clk, rst)
      begin
        if rst = '1' then
          state <= A;
        elsif rising_edge(clk) then
          state <= next_state;
        end if;
      end process;

      process(state, x)
      begin
        z <= '0';
        case state is
          when A =>
            if x = '1' then next_state <= B;
            else next_state <= A;
            end if;
          when B =>
            if x = '0' then next_state <= C;
            else next_state <= B;
            end if;
          when C =>
            if x = '0' then next_state <= D;
            else next_state <= B;
            end if;
          when D =>
            if x = '1' then
              next_state <= B;
              z <= '1';
            else
              next_state <= A;
            end if;
          when others =>
            next_state <= A;
        end case;
      end process;
    end behavior;
    """)
    add_answer_item(doc, 2, "三个 process 说明", """
    第一个 process 是状态寄存器，负责复位和时钟上升沿更新 state。
    第二个 process 是组合逻辑，既产生 next_state，也产生 Mealy 输出 z。
    本题 z 与当前状态和当前输入有关；在组合 process 开头先令 z <= '0'，可避免遗漏分支导致锁存。
    """)
    doc.save(SYNC_ANSWER_DOCX)
    return SYNC_ANSWER_DOCX


def structural_check(path):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    if "历年" in text:
        raise RuntimeError(f"{path.name} contains forbidden reference to past exam text")
    return len(doc.paragraphs), len(doc.inline_shapes), len(doc.tables)


def main():
    figs = make_figures()
    q_path = build_question_doc(figs)
    a_path = build_answer_doc(figs)
    s_path = build_sync_doc(figs)
    sa_path = build_sync_answer_doc(figs)
    for path in [q_path, a_path, s_path, sa_path]:
        paragraphs, images, tables = structural_check(path)
        print(f"{path.name}: paragraphs={paragraphs}, images={images}, tables={tables}")
    print(q_path)
    print(a_path)
    print(s_path)
    print(sa_path)


if __name__ == "__main__":
    main()
