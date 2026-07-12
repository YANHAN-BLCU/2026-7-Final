from pathlib import Path
import os
import re
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


BASE = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent
OUT.mkdir(exist_ok=True)
MD = BASE / "计算机组成原理期末复习资料.md"
DOCX = OUT / "计算机组成原理2025真题精讲补充.docx"
PDF = OUT / "计算机组成原理2025真题精讲补充.pdf"


def get_section_11():
    text = MD.read_text(encoding="utf-8")
    m = re.search(r"^## 11\. 2025 年真题精讲\s*$", text, flags=re.M)
    if not m:
        raise RuntimeError("未找到 2025 年真题精讲章节")
    return text[m.start():].strip()


def set_font(run, east="宋体", latin="Calibri", size=10.5, bold=False, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=9.5, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "center")
    tc_pr.append(v_align)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_from_md(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*-+", line):
            continue
        row = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(row)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell(table.cell(i, j), value, bold=(i == 0))
            if i == 0:
                shade_cell(table.cell(i, j), "E8EEF5")


def flush_list(doc, items, ordered=False):
    for item in items:
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        r = p.add_run(item)
        set_font(r, size=10.5)


def build_docx():
    md = get_section_11()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for name, size in [("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12)]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string("1F4D78")
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)

    lines = md.splitlines()
    table_lines = []
    code_lines = []
    in_code = False

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                for c in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    r = p.add_run(c)
                    r.font.name = "Consolas"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                    r.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                if table_lines:
                    add_table_from_md(doc, table_lines)
                    table_lines = []
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            add_table_from_md(doc, table_lines)
            table_lines = []

        if not line.strip():
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, east="黑体", size=20, bold=True, color="0B2545")
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), style="Heading 1")
        elif line.startswith("#### "):
            add_para(doc, line[5:].strip(), style="Heading 2")
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(re.sub(r"^\d+\.\s+", "", line))
            set_font(r)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:].strip())
            set_font(r)
        elif line.startswith("> "):
            p = add_para(doc, line[2:].strip(), size=10.5, color="555555")
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            add_para(doc, line)

    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.save(DOCX)
    return DOCX


def export_pdf(docx_path):
    ps = f"""
$ErrorActionPreference='Stop'
$word = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open('{str(docx_path)}', $false, $true)
  $doc.ExportAsFixedFormat('{str(PDF)}', 17)
  $doc.Close($false)
}} finally {{
  if ($word -ne $null) {{ try {{ $word.Quit() }} catch {{ }} }}
}}
"""
    subprocess.run(["powershell", "-NoProfile", "-Command", ps], check=True)
    return PDF


def build_pdf_direct():
    font_candidates = [
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]
    font_path = next((p for p in font_candidates if p.exists()), None)
    if font_path is None:
        raise RuntimeError("未找到中文字体，无法生成 PDF")
    pdfmetrics.registerFont(TTFont("CN", str(font_path)))

    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "CNNormal",
        parent=styles["Normal"],
        fontName="CN",
        fontSize=10.5,
        leading=16,
        spaceAfter=5,
    )
    h1 = ParagraphStyle(
        "CNH1",
        parent=normal,
        fontName="CN",
        fontSize=17,
        leading=22,
        textColor=colors.HexColor("#0B2545"),
        alignment=1,
        spaceBefore=8,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "CNH2",
        parent=normal,
        fontName="CN",
        fontSize=14,
        leading=19,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=10,
        spaceAfter=6,
    )
    h3 = ParagraphStyle(
        "CNH3",
        parent=normal,
        fontName="CN",
        fontSize=12,
        leading=17,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=8,
        spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "Code",
        parent=normal,
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        leftIndent=12,
    )

    def esc(s):
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("`", "")
        )

    story = []
    lines = get_section_11().splitlines()
    table_lines = []
    code_lines = []
    in_code = False

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for line in table_lines:
            if re.match(r"^\|\s*-+", line):
                continue
            rows.append([esc(c.strip()) for c in line.strip().strip("|").split("|")])
        if rows:
            t = Table(rows, repeatRows=1)
            t.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), "CN"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9AA5B1")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 6))
        table_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_lines), code_style))
                story.append(Spacer(1, 6))
                code_lines = []
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("## "):
            story.append(Paragraph(esc(line[3:].strip()), h1))
        elif line.startswith("### "):
            story.append(Paragraph(esc(line[4:].strip()), h2))
        elif line.startswith("#### "):
            story.append(Paragraph(esc(line[5:].strip()), h3))
        elif re.match(r"^\d+\.\s+", line):
            story.append(Paragraph("• " + esc(re.sub(r"^\d+\.\s+", "", line)), normal))
        elif line.startswith("- "):
            story.append(Paragraph("• " + esc(line[2:].strip()), normal))
        elif line.startswith("> "):
            story.append(Paragraph(esc(line[2:].strip()), normal))
        else:
            story.append(Paragraph(esc(line), normal))

    flush_table()
    doc = SimpleDocTemplate(
        str(PDF),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
    )
    doc.build(story)
    return PDF


if __name__ == "__main__":
    docx = build_docx()
    print(docx)
    try:
        pdf = export_pdf(docx)
        print(pdf)
    except Exception as exc:
        print(f"PDF 导出失败：{exc}")
        pdf = build_pdf_direct()
        print(pdf)
